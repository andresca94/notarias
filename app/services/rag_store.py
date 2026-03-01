from __future__ import annotations
import re
import subprocess
import tempfile
from pathlib import Path
from typing import Dict, List, Tuple

try:
    from rank_bm25 import BM25Okapi
except Exception:
    BM25Okapi = None

# Stopwords para token-overlap (palabras vacías y términos RAG genéricos)
_STOPS = {
    "de", "del", "la", "el", "los", "las", "y", "en", "a", "o", "u",
    "con", "sin", "al", "que", "minuta", "bien", "bienes",
}

# Tipos de acto conocidos para penalización BM25
_ACTO_TYPES = [
    "compraventa", "hipoteca", "cancelacion", "cambio",
    "afectacion", "dacion", "insinuacion", "aclaracion",
    "donacion", "actualizacion", "derechos",
]


# ── Helpers de extracción para Word docs ────────────────────────────────────

def _extract_docx(path: Path) -> str:
    """Extrae texto de un .docx usando python-docx."""
    from docx import Document  # type: ignore
    doc = Document(str(path))
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text)
    return "\n".join(parts)


def _extract_doc(path: Path) -> str:
    """
    .doc → texto plano.
    Intenta LibreOffice primero (Linux/CI); cae a textutil (macOS) si no está.
    """
    import shutil

    def _via_libreoffice() -> str:
        with tempfile.TemporaryDirectory() as tmpdir:
            subprocess.run(
                ["libreoffice", "--headless", "--convert-to", "txt:Text",
                 "--outdir", tmpdir, str(path)],
                capture_output=True,
                timeout=30,
            )
            out = Path(tmpdir) / (path.stem + ".txt")
            return out.read_text(encoding="utf-8", errors="ignore") if out.exists() else ""

    def _via_textutil() -> str:
        result = subprocess.run(
            ["textutil", "-convert", "txt", "-stdout", str(path)],
            capture_output=True,
            timeout=30,
        )
        return result.stdout.decode("utf-8", errors="ignore") if result.returncode == 0 else ""

    if shutil.which("libreoffice"):
        return _via_libreoffice()
    if shutil.which("textutil"):
        return _via_textutil()
    return ""


def _strip_formulario_header(text: str) -> str:
    """
    Elimina el formulario de calificación notarial que encabeza cada minuta Word.
    Conserva solo el texto narrativo que comienza con 'Compareció', 'Comparecieron',
    'En la ciudad de', 'En el municipio de', etc.
    NO usa ancla ^ para compatibilidad con minutas donde "En la ciudad de" aparece
    inline (concatenado con guiones en el mismo párrafo, no al inicio de línea).
    """
    m = re.search(
        r'(Compareció\b|Comparecieron\b|En la ciudad de\b|En el municipio de\b)',
        text,
        re.IGNORECASE,
    )
    return text[m.start():].strip() if m else text


def _strip_boilerplate_tail(text: str) -> str:
    """
    Elimina el bloque OTORGAMIENTO Y AUTORIZACIÓN (boilerplate idéntico en todas las
    minutas) y todo lo que le sigue: PROTECCIÓN DE DATOS, NOTIFICACIONES ELECTRÓNICAS,
    INSERTOS, PAZ Y SALVOS, UIAF, FIRMAS.
    El contenido específico del acto termina justo antes de ese bloque.
    """
    # El patrón varía entre minutas pero siempre tiene "O T O R G A M I E N T O"
    # rodeado de guiones/espacios
    m = re.search(
        r'[-\s]{3,}O\s*T\s*O\s*R\s*G\s*A\s*M\s*I\s*E\s*N\s*T\s*O',
        text,
        re.IGNORECASE,
    )
    if m:
        return text[:m.start()].strip()
    # Fallback: cortar en "INSERTOS:" si el bloque de otorgamiento no se encontró
    m2 = re.search(r'^\s*INSERTOS\s*:', text, re.MULTILINE | re.IGNORECASE)
    if m2:
        return text[:m2.start()].strip()
    return text


def _wrap_as_template(stem: str, text: str) -> str:
    """Envuelve texto extraído con el header estándar de plantilla RAG."""
    sep = "=" * 78
    return f"{sep}\nACTO: {stem.upper()}\n{sep}\n{{Cuerpo del Acto}}\n{text}"


# ── LocalRAGStore ────────────────────────────────────────────────────────────

class LocalRAGStore:
    """
    RAG local: carga .txt de rag_store (prioridad alta) y minutas .doc/.docx
    de rag_store/actos/ (prioridad base). Recupera el más relevante por BM25.
    """

    def __init__(self, rag_dir: str):
        self.dir = Path(rag_dir)
        self.docs: List[Tuple[str, str]] = []  # (filename, text)
        self._priorities: List[float] = []     # multiplicador de score por doc
        self._bm25 = None
        self._tokens = None

    def load(self) -> None:
        self.docs = []
        self._priorities = []

        actos_dir = self.dir / "actos"
        if not actos_dir.is_dir():
            return

        # Stems que ya tienen versión .txt limpia (prioridad sobre Word)
        txt_stems = {p.stem.lower() for p in actos_dir.iterdir() if p.suffix.lower() == ".txt"}

        for p in sorted(actos_dir.iterdir()):
            suffix = p.suffix.lower()
            try:
                if suffix == ".txt":
                    # .txt = versión pre-limpiada y editable manualmente
                    raw = p.read_text(encoding="utf-8", errors="ignore")
                    raw = _strip_boilerplate_tail(raw)  # por si acaso aún tiene el bloque
                elif suffix in (".docx", ".doc"):
                    # Saltar si ya existe un .txt para este stem
                    if p.stem.lower() in txt_stems:
                        continue
                    raw = _extract_docx(p) if suffix == ".docx" else _extract_doc(p)
                    raw = _strip_formulario_header(raw)
                    raw = _strip_boilerplate_tail(raw)
                else:
                    continue
                if raw.strip():
                    self.docs.append((p.name, _wrap_as_template(p.stem, raw)))
                    self._priorities.append(1.0)
            except Exception:
                continue  # skip silently

        if BM25Okapi and self.docs:
            self._tokens = [self._tokenize(t) for _, t in self.docs]
            self._bm25 = BM25Okapi(self._tokens)

    def _tokenize(self, s: str) -> List[str]:
        s = s.lower()
        for ch in ",.;:()[]{}<>/\\\n\r\t":
            s = s.replace(ch, " ")
        return [w for w in s.split(" ") if w.strip()]

    def _kw_score(self, q_keywords: List[str], fn_norm: str) -> float:
        """
        Fracción de keywords del query encontrados en el filename normalizado.
        Soporta variación singular/plural: compra↔compras, inmueble↔inmuebles.
        """
        if not q_keywords:
            return 0.0
        fn_tok = set(fn_norm.split())
        # Expandir con variantes singular/plural
        fn_exp: set = set()
        for t in fn_tok:
            fn_exp.add(t)
            if t.endswith("s") and len(t) > 3:
                fn_exp.add(t[:-1])   # inmuebles → inmueble
            else:
                fn_exp.add(t + "s")  # inmueble → inmuebles
        hit = 0
        for kw in q_keywords:
            if kw in fn_exp:
                hit += 1
            elif kw.endswith("s") and kw[:-1] in fn_exp:
                hit += 1
            elif not kw.endswith("s") and kw + "s" in fn_exp:
                hit += 1
        return hit / len(q_keywords)

    def search_acto(self, acto_query: str, top_k: int = 1) -> List[Dict]:
        if not self.docs:
            self.load()
        if not self.docs:
            return []

        q = acto_query.strip().lower()
        q_norm = q.replace("_", " ").replace("-", " ").replace(".", " ")

        # ── A1: Token-overlap con normalización plural/singular ────────────────
        # Elimina stopwords y términos RAG genéricos; busca keywords reales en filename
        q_keywords = [t for t in q_norm.split() if t not in _STOPS and len(t) > 2]

        best_score = 0.0
        best_exact: List[Dict] = []
        for fn, txt in self.docs:
            fn_norm2 = fn.lower().replace("_", " ").replace("-", " ").replace(".", " ")
            sc = self._kw_score(q_keywords, fn_norm2)
            if sc > best_score:
                best_score = sc
                best_exact = [{"file": fn, "text": txt, "score": 999.0 * sc}]
            elif sc == best_score and sc > 0:
                best_exact.append({"file": fn, "text": txt, "score": 999.0 * sc})

        if best_score >= 0.5:
            return best_exact[:top_k]

        # ── A2: BM25 fallback con penalización de tipo de acto incorrecto ──────
        if self._bm25:
            qtok = self._tokenize(q)
            scores = self._bm25.get_scores(qtok)
            # Detectar el tipo de acto en el query
            q_tipo = next((t for t in _ACTO_TYPES if t in q_norm), None)
            ranked_list = []
            for idx, sc in enumerate(scores):
                fn_lower = self.docs[idx][0].lower()
                adj = float(sc) * self._priorities[idx]
                # Penalizar fuertemente si el tipo de acto no coincide
                if q_tipo and q_tipo not in fn_lower:
                    adj *= 0.05
                ranked_list.append((idx, adj))
            ranked_list.sort(key=lambda x: x[1], reverse=True)
            out = []
            for idx, adj in ranked_list[:top_k]:
                fn, txt = self.docs[idx]
                out.append({"file": fn, "text": txt, "score": adj})
            return out

        # fallback: substring
        out = []
        for fn, txt in self.docs:
            sc = 1.0 if q in txt.lower() else 0.0
            out.append({"file": fn, "text": txt, "score": sc})
        out.sort(key=lambda d: d["score"], reverse=True)
        return out[:top_k]
