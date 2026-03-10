# app/services/rendering/docx_renderer.py
from __future__ import annotations

import re
from pathlib import Path
from typing import Dict, List, Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor


# ── Colores notariales ──────────────────────────────────────────────────────
_COLOR_PENDIENTE = RGBColor(0xCC, 0x00, 0x00)  # Rojo  → [[PENDIENTE: ...]]
_COLOR_VARIABLE  = RGBColor(0xE6, 0x4A, 0x00)  # Naranja → [[VARIABLE_LLENA]]
_COLOR_GRIS      = RGBColor(0xAA, 0xAA, 0xAA)  # Gris → separadores de guiones

# ── Patrones regex ───────────────────────────────────────────────────────────
# Variables notariales [[...]]
_VAR_RE = re.compile(r'(\[\[.*?\]\])', re.DOTALL)

# Secuencias de 2+ palabras totalmente en MAYÚSCULAS (nombres/entidades)
_CAPS_SEQ = re.compile(
    r'[A-ZÁÉÍÓÚÜÑ]{2,}(?:\s[A-ZÁÉÍÓÚÜÑ]{2,})+'
)

# Línea de encabezado: "CAMPO: valor"
_HEADER_LINE_RE = re.compile(r'^[A-ZÁÉÍÓÚÜÑ][A-ZÁÉÍÓÚÜÑ0-9 \(\)\/]{1,50}:\s')

# Título de acto EP: ---PRIMER ACTO--- o ---CANCELACION PACTO DE RETROVENTA---
_ACT_TITLE_RE = re.compile(r'^-{2,}\s*([A-ZÁÉÍÓÚÜÑ0-9][A-ZÁÉÍÓÚÜÑ0-9 ÁÉÍÓÚ]*[A-ZÁÉÍÓÚÜÑ0-9])\s*-{2,}$')

# Número de cláusula al inicio de bloque (PRIMERO., SEGUNDO:, PARÁGRAFO PRIMERO:, etc.)
_CLAUSE_NUM_RE = re.compile(
    r'^(PRIMERO|SEGUNDO|TERCERO|CUARTO|QUINTO|SEXTO|SÉPTIMO|OCTAVO|NOVENO|DÉCIMO'
    r'|PARÁGRAFO\s+(?:PRIMERO|SEGUNDO|TERCERO|CUARTO|ÚNICO)'
    r'|CLÁUSULA\s+\w+)[\.\:\-\s]',
    re.IGNORECASE
)

# Separador puro de guiones (≥10 guiones, solo guiones)
_DASH_SEP_RE = re.compile(r'^[-─]{10,}$')

# Guiones al FINAL de un bloque (para extender a relleno de línea completa)
_TRAILING_DASH_RE = re.compile(r'(-{3,})\s*\t?\s*$', re.MULTILINE)

# Guiones de relleno estándar al estilo EP colombiano
_LINE_FILL_DASHES = '-' * 90


# ── Helpers de detección ─────────────────────────────────────────────────────

def _is_header_block(block: str) -> bool:
    """True si ≥60% de las líneas son del tipo 'CAMPO: valor'."""
    lines = [l.strip() for l in block.split('\n') if l.strip()]
    if len(lines) < 2:
        return False
    n = sum(1 for l in lines if _HEADER_LINE_RE.match(l))
    return n >= len(lines) * 0.6


def _is_title_line(line: str) -> bool:
    """True si la línea es un título de sección predominantemente en mayúsculas.
    Ejemplos: 'OTORGAMIENTO y AUTORIZACIÓN', 'DE LA CAPACIDAD:', 'DERECHOS NOTARIALES:'
    """
    s = line.strip()
    if not s or len(s) > 100:
        return False
    letters = [c for c in s if c.isalpha()]
    if len(letters) < 3:
        return False
    upper_ratio = sum(1 for c in letters if c.isupper()) / len(letters)
    return upper_ratio >= 0.65 and len(s.split()) <= 12  # ampliado de 10 → 12


def _is_act_title_line(line: str) -> bool:
    """True si la línea es un título de acto EP: ---PRIMER ACTO--- """
    return bool(_ACT_TITLE_RE.match(line.strip()))


def _is_act_title_block(block: str) -> bool:
    """True si TODAS las líneas no vacías del bloque son títulos de acto."""
    lines = [l for l in block.split('\n') if l.strip()]
    return bool(lines) and all(_is_act_title_line(l) for l in lines)


def _is_dash_sep_block(block: str) -> bool:
    """True si el bloque es exclusivamente guiones (separador visual)."""
    lines = [l.strip() for l in block.split('\n') if l.strip()]
    return bool(lines) and all(_DASH_SEP_RE.match(l) for l in lines)


def _preprocess_block(block: str) -> str:
    """Extiende secuencias de guiones finales para completar la línea visual al estilo EP."""
    return _TRAILING_DASH_RE.sub(_LINE_FILL_DASHES, block)


# ── Helpers de formato ───────────────────────────────────────────────────────

def _add_line_to_para(para, line: str) -> None:
    """
    Agrega una línea al párrafo con:
    - Variables [[...]] coloreadas (rojo para PENDIENTE, naranja para llenas).
    - Secuencias en MAYÚSCULAS (nombres / entidades) en negrita.
    - Resto del texto en formato normal.
    """
    segments = _VAR_RE.split(line)
    for seg in segments:
        if not seg:
            continue
        if re.fullmatch(r'\[\[.*?\]\]', seg, re.DOTALL):
            run = para.add_run(seg)
            run.bold = True
            if "PENDIENTE" in seg.upper():
                run.font.color.rgb = _COLOR_PENDIENTE
            else:
                run.font.color.rgb = _COLOR_VARIABLE
        else:
            pos = 0
            for m in _CAPS_SEQ.finditer(seg):
                before = seg[pos:m.start()]
                if before:
                    para.add_run(before)
                run = para.add_run(m.group())
                run.bold = True
                pos = m.end()
            tail = seg[pos:]
            if tail:
                para.add_run(tail)


def _add_header_line_to_para(para, line: str) -> None:
    """Como _add_line_to_para pero también bold en el label 'CAMPO:' de líneas de encabezado.
    FIX v40-MMFL1: _CAPS_SEQ requiere 2+ palabras → 'RADICADO', 'FECHA', 'VALOR' (1 palabra) no se boldan.
    Esta función bold el label explícitamente para encabezados tipo 'CAMPO: valor'.
    """
    m = _HEADER_LINE_RE.match(line)
    if m:
        label = line[:m.end()].rstrip()  # e.g. "RADICADO:" o "OTORGANTE(S):"
        rest = line[m.end():]
        run = para.add_run(label)
        run.bold = True
        if rest:
            _add_line_to_para(para, rest)
    else:
        _add_line_to_para(para, line)


def _set_spacing(para, before_pt: float = 0, after_pt: float = 0) -> None:
    """Aplica espaciado antes/después a un párrafo."""
    fmt = para.paragraph_format
    if before_pt:
        fmt.space_before = Pt(before_pt)
    if after_pt:
        fmt.space_after = Pt(after_pt)


def _insert_table_after(paragraph, text_block: str) -> None:
    """
    Crea una tabla de una columna donde cada línea del bloque es una fila.
    La tabla se inserta inmediatamente después de 'paragraph'.
    """
    lines = [
        l for l in text_block.split("\n")
        if l.strip() not in ("###TABLE_START###", "###TABLE_END###")
    ]
    if not lines:
        return

    parent = paragraph._p.getparent()
    idx = list(parent).index(paragraph._p)

    doc = paragraph.part.document
    table = doc.add_table(rows=len(lines), cols=1)
    try:
        table.style = "Table Grid"
    except Exception:
        pass

    for i, row_text in enumerate(lines):
        cell = table.rows[i].cells[0]
        cell.text = ""
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        if _is_title_line(row_text):
            run = para.add_run(row_text.strip())
            run.bold = True
        else:
            _add_line_to_para(para, row_text)

    tbl_xml = table._tbl
    parent.remove(tbl_xml)
    parent.insert(idx + 1, tbl_xml)


def _render_block(b: str, doc, parent, idx: int) -> int:
    """
    Renderiza un bloque de texto como uno o más párrafos Word con formato notarial.
    Devuelve el nuevo índice de inserción.

    Jerarquía de tipos:
    1. Tabla UIAF (###TABLE_START###)
    2. Separador de guiones puro
    3. Título de acto EP (---PRIMER ACTO---)
    4. Bloque de encabezado CAMPO: valor
    5. Cláusula numerada (PRIMERO., SEGUNDO:, etc.)
    6. Línea de título corta (OTORGAMIENTO, DE LA CAPACIDAD:, etc.)
    7. Texto legal justificado (default)
    """
    # 1) Tabla UIAF
    if b.strip().startswith("###TABLE_START###"):
        anchor = doc.add_paragraph()
        anchor_xml = anchor._p
        parent.remove(anchor_xml)
        parent.insert(idx + 1, anchor_xml)
        idx += 1
        _insert_table_after(anchor, b)
        idx += 1
        return idx

    # Fix 3 v33: Limpiar marcadores TABLE que aparecen INLINE (no detectados como bloques tabla)
    # Causa: DataBinder o binder concatena marcadores en la misma línea que contenido real
    # Ejemplo del artifact PDF-32: "DIRECCIÓN: ###TABLETABLE_START### CIUDAD: ###TABLE_END### EMAIL:..."
    if "###TABLE" in b:
        b = re.sub(r'###TABLE[A-Z_]*###', '', b).strip()
        if not b:
            return idx

    # Pre-procesado: extender guiones finales
    b = _preprocess_block(b)

    # 2) Separador puro de guiones
    if _is_dash_sep_block(b):
        new_p = doc.add_paragraph()
        new_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = new_p.add_run(_LINE_FILL_DASHES)
        run.font.color.rgb = _COLOR_GRIS
        _set_spacing(new_p, after_pt=0)
        p_xml = new_p._p
        parent.remove(p_xml)
        parent.insert(idx + 1, p_xml)
        idx += 1
        return idx

    # 3) Título de acto EP (---PRIMER ACTO--- / ---CANCELACION PACTO DE RETROVENTA---)
    if _is_act_title_block(b):
        lines = [l for l in b.split('\n') if l.strip()]
        first = True
        for line in lines:
            m = _ACT_TITLE_RE.match(line.strip())
            text = m.group(1).strip() if m else line.strip('- ').strip()
            new_p = doc.add_paragraph()
            new_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = new_p.add_run(text)
            run.bold = True
            run.font.size = Pt(13)
            if first:
                _set_spacing(new_p, before_pt=18, after_pt=2)
                first = False
            else:
                _set_spacing(new_p, after_pt=6)
            p_xml = new_p._p
            parent.remove(p_xml)
            parent.insert(idx + 1, p_xml)
            idx += 1
        return idx

    # 4) Encabezado CAMPO: valor
    if _is_header_block(b):
        new_p = doc.add_paragraph()
        new_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        lines = b.split('\n')
        for i, line in enumerate(lines):
            if i > 0:
                br_run = new_p.add_run()
                br_run.add_break()
            _add_header_line_to_para(new_p, line)  # FIX v40-MMFL1: bold labels como RADICADO:, FECHA:
        p_xml = new_p._p
        parent.remove(p_xml)
        parent.insert(idx + 1, p_xml)
        idx += 1
        return idx

    # 5) Cláusula numerada (PRIMERO., SEGUNDO:, PARÁGRAFO PRIMERO:, etc.)
    first_line = b.split('\n')[0].strip()
    clause_match = _CLAUSE_NUM_RE.match(first_line)
    if clause_match:
        new_p = doc.add_paragraph()
        new_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        _set_spacing(new_p, before_pt=4, after_pt=2)
        lines = b.split('\n')
        for i, line in enumerate(lines):
            if i > 0:
                br_run = new_p.add_run()
                br_run.add_break()
            if i == 0:
                # Separar el número de cláusula del resto
                cm = _CLAUSE_NUM_RE.match(line.strip())
                if cm:
                    label_end = cm.end()
                    label = line[:label_end].rstrip()
                    rest = line[label_end:]
                    # E-03: Extender bold al subtítulo all-caps hasta el primer ":"
                    # Ejemplo: "TERCERO-" + " LIBERTAD: ..." → bold "TERCERO- LIBERTAD:"
                    _rest_s = rest.lstrip('- \t')
                    _col = _rest_s[:60].find(':')
                    if _col >= 0:
                        _sub = _rest_s[:_col].strip()
                        if _sub and re.match(r'^[A-ZÁÉÍÓÚÜÑ\s]+$', _sub):
                            _lead = rest[:len(rest) - len(_rest_s)]
                            run = new_p.add_run(label + _lead + _rest_s[:_col + 1])
                            run.bold = True
                            run.font.bold = True  # Fix-3 v13: forzar a nivel font para override de estilo
                            rest = _rest_s[_col + 1:]
                        else:
                            run = new_p.add_run(label)
                            run.bold = True
                            run.font.bold = True  # Fix-3 v13
                    else:
                        run = new_p.add_run(label)
                        run.bold = True
                        run.font.bold = True  # Fix-3 v13
                    if rest:
                        _add_line_to_para(new_p, rest)
                else:
                    _add_line_to_para(new_p, line)
            else:
                _add_line_to_para(new_p, line)
        p_xml = new_p._p
        parent.remove(p_xml)
        parent.insert(idx + 1, p_xml)
        idx += 1
        return idx

    # 6 + 7) Líneas de título y texto legal (procesado línea a línea)
    new_p = doc.add_paragraph()
    is_header = False  # ya descartado arriba
    new_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    _set_spacing(new_p, after_pt=2)

    lines = b.split('\n')
    for i, line in enumerate(lines):
        if i > 0:
            br_run = new_p.add_run()
            br_run.add_break()
        if _is_title_line(line):
            run = new_p.add_run(line.strip())
            run.bold = True
            # Si es la única/primera línea del bloque, añadir espacio
            if i == 0:
                _set_spacing(new_p, before_pt=8, after_pt=2)
        else:
            _add_line_to_para(new_p, line)

    p_xml = new_p._p
    parent.remove(p_xml)
    parent.insert(idx + 1, p_xml)
    idx += 1
    return idx


def _insert_paragraphs_after(paragraph, blocks: List[str]) -> None:
    """
    Inserta bloques como párrafos notariales después de 'paragraph'.
    Aplica formato visual EP colombiano:
    - Títulos de acto (---PRIMER ACTO---): bold, centrado, 13pt, espaciado
    - Cláusulas numeradas (PRIMERO., SEGUNDO:): número en negrita
    - Separadores de guiones: extendidos a 90 chars, color gris
    - Guiones al final de párrafo: extendidos a 90 chars
    - Títulos de sección cortos: negrita
    - Encabezados CAMPO: valor: alineación izquierda
    - Texto legal: justificado
    """
    parent = paragraph._p.getparent()
    idx = parent.index(paragraph._p)
    doc = paragraph.part.document

    for b in blocks:
        idx = _render_block(b, doc, parent, idx)


def _replace_in_paragraph(paragraph, key: str, value: str) -> None:
    """Reemplaza texto aunque el placeholder esté partido en runs."""
    if key not in paragraph.text:
        return

    full = "".join(run.text for run in paragraph.runs)
    if key not in full:
        return

    full = full.replace(key, value)

    for run in paragraph.runs:
        run.text = ""
    paragraph.runs[0].text = full


def render_docx(template_path: str, context: Dict, out_docx_path: str) -> str:
    """
    Render DOCX "notarial-friendly".
    - Reemplaza placeholders simples tipo {{RADICADO}} en todo el documento.
    - Para {{CONTENIDO_IA}}: inserta párrafos con formato notarial EP colombiano.
    """
    template_path = str(template_path)
    out_docx_path = str(out_docx_path)

    doc = Document(template_path)

    # 1) Placeholders simples ({{RADICADO}}, {{COMENTARIO}}, etc.)
    simple_keys = []
    for k, v in context.items():
        if k == "CONTENIDO_IA":
            continue
        simple_keys.append((f"{{{{{k}}}}}", "" if v is None else str(v)))
        simple_keys.append((f"{{{{ {k} }}}}", "" if v is None else str(v)))

    # 2) Contenido IA → párrafos
    contenido = context.get("CONTENIDO_IA") or ""
    contenido = str(contenido).replace("\r\n", "\n").replace("\r", "\n").strip()
    # Fix 5 v32: NFC normalize + regex \bDECONOCIMIENTO\b (cubre Unicode decomposed y variantes de spacing)
    import unicodedata as _ud_rend
    contenido = _ud_rend.normalize('NFC', contenido)
    # Fix 2 v33: sin \b — captura DECONOCIMIENTO incluso pegado a otro carácter (edge cases encoding)
    contenido = re.sub(r'DECONOCIMIENTO', 'DE CONOCIMIENTO', contenido, flags=re.IGNORECASE)
    # A-02 adicional: cobertura extra para "CLAUSULA DECONOCIMIENTO" como frase completa
    contenido = re.sub(r'CL[ÁA]USULA\s+DECONOCIMIENTO', 'CLÁUSULA DE CONOCIMIENTO', contenido, flags=re.IGNORECASE)
    # FIX v40-MMFL1: Garantizar bloque separado para cada cláusula numerada.
    # DataBinder a veces no inserta línea en blanco entre cláusulas → todo queda en un bloque →
    # solo la primera línea se evalúa como clause_match → TERCERO/CUARTO/QUINTO/SEXTO sin bold.
    _CLAUSE_SEP_RE = re.compile(
        r'(?<!\n)\n((?:PRIMERO|SEGUNDO|TERCERO|CUARTO|QUINTO|SEXTO|S[EÉ]PTIMO|OCTAVO|NOVENO|D[EÉ]CIMO'
        r'|PAR[ÁA]GRAFO\s+(?:PRIMERO|SEGUNDO|TERCERO|CUARTO|[ÚU]NICO)'
        r'|CL[ÁA]USULA\s+\w+)[\.\:\-\s])',
        re.IGNORECASE
    )
    contenido = _CLAUSE_SEP_RE.sub(r'\n\n\1', contenido)
    contenido = re.sub(r'\n{3,}', '\n\n', contenido)  # limpiar triple saltos
    blocks = [b.strip() for b in contenido.split("\n\n") if b.strip()]

    for p in doc.paragraphs:
        for ph, val in simple_keys:
            if ph in p.text:
                _replace_in_paragraph(p, ph, val)
        # A-01: Fix "CLÁUSULA DECONOCIMIENTO" hardcoded in template DOCX paragraphs
        if "DECONOCIMIENTO" in p.text.upper():
            _replace_in_paragraph(p, "DECONOCIMIENTO", "DE CONOCIMIENTO")

    # FIX 3 v39: iterar w:t XML directamente — funciona con y sin w:r wrapper.
    # Fix E v38 usaba _tp_e.runs (solo w:r), pero si el texto está en w:t directo sin w:r,
    # _tp_e.runs=[] → _full_e="" → no reemplazaba nada. w:t.iter() alcanza todos los nodos de texto.
    _W_T_NS = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t'
    for _tbl_3 in doc.tables:
        for _trow_3 in _tbl_3.rows:
            for _tcell_3 in _trow_3.cells:
                if "DECONOCIMIENTO" not in _tcell_3.text.upper():
                    continue
                for _wt_3 in _tcell_3._tc.iter(_W_T_NS):
                    if _wt_3.text and "DECONOCIMIENTO" in _wt_3.text.upper():
                        _wt_3.text = re.sub(
                            r'DECONOCIMIENTO', 'DE CONOCIMIENTO',
                            _wt_3.text, flags=re.IGNORECASE
                        )
    # FIX 3b v39: mismo fix para párrafos regulares (complementa A-01 que usa _replace_in_paragraph)
    for _p_3 in doc.paragraphs:
        if "DECONOCIMIENTO" in _p_3.text.upper():
            for _wt_3 in _p_3._p.iter(_W_T_NS):
                if _wt_3.text and "DECONOCIMIENTO" in _wt_3.text.upper():
                    _wt_3.text = re.sub(
                        r'DECONOCIMIENTO', 'DE CONOCIMIENTO',
                        _wt_3.text, flags=re.IGNORECASE
                    )

    # Busca el párrafo donde está el placeholder de contenido
    placeholder_found = False
    for p in doc.paragraphs:
        if "{{CONTENIDO_IA}}" in p.text or "{{ CONTENIDO_IA }}" in p.text:
            placeholder_found = True
            _replace_in_paragraph(p, "{{CONTENIDO_IA}}", "")
            _replace_in_paragraph(p, "{{ CONTENIDO_IA }}", "")
            _insert_paragraphs_after(p, blocks)
            break

    if not placeholder_found:
        doc.add_paragraph("")
        parent = doc.element.body
        idx = len(list(parent)) - 1
        for b in blocks:
            idx = _render_block(b, doc, parent, idx)

    Path(out_docx_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_docx_path)
    return out_docx_path
