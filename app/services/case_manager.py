from __future__ import annotations

import difflib
import json
import os
import re
import shutil
import uuid
from contextlib import contextmanager
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Dict, Iterator, List, Optional, Tuple

from fastapi import UploadFile

from app.core.config import settings


DOCX_MIME_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
PDF_MIME_TYPE = "application/pdf"
TEXT_MARKDOWN_MIME_TYPE = "text/markdown; charset=utf-8"
CASE_FILENAME_HINT_RE = re.compile(r"(?:caso|radicado)[\s_-]*(\d{4,})", re.IGNORECASE)


class CaseStateError(RuntimeError):
    pass


class CaseLockError(RuntimeError):
    pass


def utc_now_iso() -> str:
    return datetime.now(timezone.utc).replace(microsecond=0).isoformat().replace("+00:00", "Z")


def _sanitize_filename(name: str) -> str:
    base = (name or "archivo").strip()
    base = base.replace("/", "_").replace("\\", "_")
    base = re.sub(r"\s+", " ", base)
    return base or "archivo"


def _unique_destination(directory: Path, original_name: str) -> Path:
    directory.mkdir(parents=True, exist_ok=True)
    original = Path(_sanitize_filename(original_name))
    stem = original.stem or "archivo"
    suffix = original.suffix
    candidate = directory / f"{stem}{suffix}"
    idx = 2
    while candidate.exists():
        candidate = directory / f"{stem}_{idx}{suffix}"
        idx += 1
    return candidate


def extract_case_hint_from_filename(filename: str) -> Optional[str]:
    match = CASE_FILENAME_HINT_RE.search((filename or "").strip())
    if not match:
        return None
    return match.group(1)


def case_dir(radicado: str) -> Path:
    return Path(settings.OUTPUT_DIR) / f"CASE-{radicado}"


def case_state_path(radicado: str) -> Path:
    return case_dir(radicado) / "case_state.json"


def incoming_root() -> Path:
    return Path(settings.OUTPUT_DIR) / "_incoming"


def feedback_corpus_root() -> Path:
    return Path(settings.OUTPUT_DIR) / "_feedback_corpus"


def feedback_corpus_events_path() -> Path:
    return feedback_corpus_root() / "feedback_events.jsonl"


def feedback_corpus_runs_path() -> Path:
    return feedback_corpus_root() / "openclaw_runs.jsonl"


def backend_maintenance_state_path() -> Path:
    return feedback_corpus_root() / "backend_maintenance_state.json"


def _parse_utc_iso(value: Optional[str]) -> Optional[datetime]:
    if not value:
        return None
    normalized = value.strip()
    if not normalized:
        return None
    try:
        return datetime.fromisoformat(normalized.replace("Z", "+00:00"))
    except ValueError:
        return None


def load_backend_maintenance_state() -> Optional[Dict[str, Any]]:
    path = backend_maintenance_state_path()
    if not path.exists():
        return None
    return json.loads(path.read_text(encoding="utf-8"))


def save_backend_maintenance_state(state: Dict[str, Any]) -> Dict[str, Any]:
    path = backend_maintenance_state_path()
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(json.dumps(state, ensure_ascii=False, indent=2), encoding="utf-8")
    return state


def refresh_backend_maintenance_state() -> Optional[Dict[str, Any]]:
    state = load_backend_maintenance_state()
    if not state:
        return None

    status = (state.get("status") or "").strip().lower()
    if status not in {"queued", "running"}:
        return state

    timeout_seconds = int(settings.OPENCLAW_MAINTENANCE_PENDING_TIMEOUT_SECONDS or 0)
    if timeout_seconds <= 0:
        return state

    anchor = _parse_utc_iso(state.get("started_at")) or _parse_utc_iso(state.get("queued_at"))
    if not anchor:
        return state

    elapsed_seconds = (datetime.now(timezone.utc) - anchor).total_seconds()
    if elapsed_seconds < timeout_seconds:
        return state

    expired_state = dict(state)
    expired_state["status"] = "failed"
    expired_state["message"] = (
        "La actualización automática del backend excedió el tiempo máximo de espera. "
        "Revisa logs y vuelve a intentar cuando el backend esté estable."
    )
    expired_state.setdefault("started_at", state.get("queued_at") or utc_now_iso())
    expired_state["finished_at"] = utc_now_iso()
    save_backend_maintenance_state(expired_state)
    return expired_state


def set_backend_maintenance_state(
    *,
    status: str,
    radicado: Optional[str] = None,
    iteration: Optional[int] = None,
    message: Optional[str] = None,
    run_id: Optional[str] = None,
) -> Dict[str, Any]:
    now = utc_now_iso()
    current = refresh_backend_maintenance_state() or {}
    state = dict(current)
    state["status"] = status
    if radicado is not None:
        state["radicado"] = str(radicado)
    if iteration is not None:
        state["iteration"] = int(iteration)
    if status == "queued":
        state["queued_at"] = now
        state.pop("started_at", None)
        state.pop("finished_at", None)
        state.pop("run_id", None)
    elif status == "running":
        state.setdefault("queued_at", now)
        state["started_at"] = now
        state.pop("finished_at", None)
    else:
        state.setdefault("queued_at", now)
        state.setdefault("started_at", now)
        state["finished_at"] = now
    if message is not None:
        state["message"] = message
    if run_id is not None:
        state["run_id"] = run_id
    return save_backend_maintenance_state(state)


def backend_maintenance_pending() -> Optional[Dict[str, Any]]:
    state = refresh_backend_maintenance_state()
    if not state:
        return None
    if (state.get("status") or "").strip().lower() in {"queued", "running"}:
        return state
    return None


def create_staging_dir() -> Path:
    path = incoming_root() / uuid.uuid4().hex
    path.mkdir(parents=True, exist_ok=True)
    return path


def relative_to_case(radicado: str, path: Path) -> str:
    return str(path.relative_to(case_dir(radicado)))


def resolve_case_path(radicado: str, relative_path: str) -> Path:
    return case_dir(radicado) / relative_path


async def save_upload_file(file: UploadFile, directory: Path) -> Path:
    destination = _unique_destination(directory, file.filename or "archivo")
    content = await file.read()
    destination.write_bytes(content)
    await file.close()
    return destination


async def save_uploads(files: List[UploadFile], directory: Path) -> List[Path]:
    saved: List[Path] = []
    for file in files:
        saved.append(await save_upload_file(file, directory))
    return saved


def list_case_inputs(radicado: str) -> Tuple[List[str], List[str]]:
    base = case_dir(radicado) / "inputs"
    scanner_dir = base / "scanner"
    docs_dir = base / "documentos"
    scan_paths = sorted(str(p) for p in scanner_dir.iterdir() if p.is_file()) if scanner_dir.exists() else []
    doc_paths = sorted(str(p) for p in docs_dir.iterdir() if p.is_file()) if docs_dir.exists() else []
    return scan_paths, doc_paths


def build_case_state(radicado: str, comentario: str, template_id: Optional[str]) -> Dict[str, Any]:
    now = utc_now_iso()
    return {
        "schema_version": 1,
        "radicado": radicado,
        "status": "generated",
        "latest_iteration": 0,
        "template_id": template_id,
        "base_comentario": comentario or "(Sin comentarios)",
        "created_at": now,
        "updated_at": now,
        "inputs": {
            "scanner_files": [],
            "documentos_files": [],
        },
        "iterations": {},
    }


def load_case_state(radicado: str) -> Dict[str, Any]:
    path = case_state_path(radicado)
    if not path.exists():
        raise CaseStateError(f"No existe estado para el radicado {radicado}.")
    return json.loads(path.read_text(encoding="utf-8"))


def save_case_state(state: Dict[str, Any]) -> None:
    radicado = str(state["radicado"])
    path = case_state_path(radicado)
    path.parent.mkdir(parents=True, exist_ok=True)
    state["updated_at"] = utc_now_iso()
    path.write_text(json.dumps(state, ensure_ascii=False, indent=2), encoding="utf-8")


def ensure_case_state(radicado: str, comentario: str, template_id: Optional[str]) -> Dict[str, Any]:
    try:
        state = load_case_state(radicado)
    except CaseStateError:
        state = build_case_state(radicado, comentario, template_id)
    state["template_id"] = template_id or state.get("template_id")
    if comentario and comentario.strip() and not (state.get("base_comentario") or "").strip():
        state["base_comentario"] = comentario
    return state


def _copy_inputs_to_case(
    radicado: str,
    staged_scans: List[Path],
    staged_docs: List[Path],
) -> Dict[str, List[Dict[str, str]]]:
    case_inputs = case_dir(radicado) / "inputs"
    scanner_dir = case_inputs / "scanner"
    docs_dir = case_inputs / "documentos"
    scanner_dir.mkdir(parents=True, exist_ok=True)
    docs_dir.mkdir(parents=True, exist_ok=True)

    saved_scans: List[Dict[str, str]] = []
    saved_docs: List[Dict[str, str]] = []

    for path in staged_scans:
        destination = _unique_destination(scanner_dir, path.name)
        shutil.copy2(path, destination)
        saved_scans.append({"name": destination.name, "relative_path": relative_to_case(radicado, destination)})

    for path in staged_docs:
        destination = _unique_destination(docs_dir, path.name)
        shutil.copy2(path, destination)
        saved_docs.append({"name": destination.name, "relative_path": relative_to_case(radicado, destination)})

    return {
        "scanner_files": saved_scans,
        "documentos_files": saved_docs,
    }


def _copy_generated_artifacts(
    radicado: str,
    iteration: int,
    result: Dict[str, Any],
) -> Dict[str, str]:
    generated_dir = case_dir(radicado) / "iterations" / str(iteration) / "generated"
    generated_dir.mkdir(parents=True, exist_ok=True)

    src_docx = Path(result["docx_path"])
    src_pdf = Path(result["pdf_path"])

    dst_docx = generated_dir / src_docx.name
    dst_pdf = generated_dir / src_pdf.name
    shutil.copy2(src_docx, dst_docx)
    shutil.copy2(src_pdf, dst_pdf)

    root_debug_dir = case_dir(radicado) / "debug"
    dst_debug_dir = generated_dir / "debug"
    if root_debug_dir.exists():
        if dst_debug_dir.exists():
            shutil.rmtree(dst_debug_dir)
        shutil.copytree(root_debug_dir, dst_debug_dir)

    return {
        "docx_path": relative_to_case(radicado, dst_docx),
        "pdf_path": relative_to_case(radicado, dst_pdf),
        "debug_dir": relative_to_case(radicado, dst_debug_dir) if dst_debug_dir.exists() else "",
    }


def _extract_docx_blocks(path: Path) -> List[str]:
    from docx import Document  # type: ignore

    doc = Document(str(path))
    blocks: List[str] = []
    seen: set[str] = set()

    def _push(text: str) -> None:
        normalized = re.sub(r"\s+", " ", (text or "").strip())
        if not normalized or normalized in seen:
            return
        seen.add(normalized)
        blocks.append(normalized)

    for paragraph in doc.paragraphs:
        _push(paragraph.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                _push(cell.text)
    return blocks


def _truncate_text(text: str, limit: int = 420) -> str:
    value = re.sub(r"\s+", " ", (text or "").strip())
    if len(value) <= limit:
        return value
    return f"{value[: limit - 1].rstrip()}…"


def _render_excerpt(blocks: List[str], *, max_items: int = 2) -> str:
    excerpt = [_truncate_text(item) for item in blocks[:max_items] if item.strip()]
    return "\n".join(f"- {item}" for item in excerpt) or "- [sin texto detectable]"


def _load_feedback_comments_from_entry(radicado: str, entry: Dict[str, Any]) -> List[Dict[str, Any]]:
    feedback = entry.get("feedback") or {}
    comments_rel = feedback.get("comments_json_path")
    if not comments_rel:
        return []
    comments_path = resolve_case_path(radicado, comments_rel)
    if not comments_path.exists():
        return []
    return json.loads(comments_path.read_text(encoding="utf-8"))


def _write_iteration_change_report(
    *,
    radicado: str,
    state: Dict[str, Any],
    iteration: int,
    artifacts: Dict[str, str],
) -> str:
    if iteration <= 1:
        return ""

    previous_entry = (state.get("iterations") or {}).get(str(iteration - 1)) or {}
    previous_docx_rel = ((previous_entry.get("artifacts") or {}).get("docx_path") or "").strip()
    current_docx_rel = (artifacts.get("docx_path") or "").strip()
    if not previous_docx_rel or not current_docx_rel:
        return ""

    previous_docx_path = resolve_case_path(radicado, previous_docx_rel)
    current_docx_path = resolve_case_path(radicado, current_docx_rel)
    if not previous_docx_path.exists() or not current_docx_path.exists():
        return ""

    try:
        previous_blocks = _extract_docx_blocks(previous_docx_path)
        current_blocks = _extract_docx_blocks(current_docx_path)
    except Exception:
        return ""

    matcher = difflib.SequenceMatcher(a=previous_blocks, b=current_blocks, autojunk=False)
    changes: List[Dict[str, Any]] = []
    counts = {"replace": 0, "insert": 0, "delete": 0}
    for tag, i1, i2, j1, j2 in matcher.get_opcodes():
        if tag == "equal":
            continue
        if tag in counts:
            counts[tag] += 1
        changes.append(
            {
                "tag": tag,
                "before": previous_blocks[i1:i2],
                "after": current_blocks[j1:j2],
            }
        )

    comments = _load_feedback_comments_from_entry(radicado, previous_entry)
    report_lines = [
        "# Reporte de cambios de iteracion",
        "",
        f"- Radicado: {radicado}",
        f"- Comparacion: iteracion {iteration - 1} -> {iteration}",
        f"- Generado: {utc_now_iso()}",
        f"- Comentarios aplicados desde la iteracion anterior: {len(comments)}",
        f"- Bloques con cambio detectados: {len(changes)}",
        f"- Reemplazos: {counts['replace']}",
        f"- Inserciones: {counts['insert']}",
        f"- Eliminaciones: {counts['delete']}",
        "",
    ]

    if comments:
        report_lines.extend(
            [
                "## Feedback previo utilizado",
                "",
            ]
        )
        for idx, item in enumerate(comments[:12], start=1):
            report_lines.extend(
                [
                    f"{idx}. {_truncate_text(item.get('comment_text') or '', 320)}",
                    f"   - Ancla: {_truncate_text(item.get('anchor_text') or '[sin ancla]', 180)}",
                ]
            )
        report_lines.append("")

    report_lines.extend(
        [
            "## Cambios detectados en el Word generado",
            "",
        ]
    )
    if not changes:
        report_lines.append("No se detectaron diferencias textuales claras entre ambas iteraciones.")
    else:
        tag_labels = {
            "replace": "Texto reemplazado",
            "insert": "Texto agregado",
            "delete": "Texto eliminado",
        }
        for idx, change in enumerate(changes[:12], start=1):
            report_lines.extend(
                [
                    f"### Cambio {idx}: {tag_labels.get(change['tag'], change['tag'])}",
                    "",
                    "**Antes**",
                    _render_excerpt(change["before"]),
                    "",
                    "**Despues**",
                    _render_excerpt(change["after"]),
                    "",
                ]
            )

    report_lines.extend(
        [
            "## Como verificar",
            "",
            "- Descarga el Word y el PDF de la iteracion actual.",
            "- Contrasta este reporte con los comentarios del Word revisado.",
            "- Si un comentario importante no aparece reflejado, vuelve a subir feedback mas especifico antes de intentar otra iteracion.",
            "",
        ]
    )

    report_path = case_dir(radicado) / "iterations" / str(iteration) / "generated" / f"Reporte_Cambios_{radicado}_iteracion_{iteration}.md"
    report_path.write_text("\n".join(report_lines), encoding="utf-8")
    return relative_to_case(radicado, report_path)


def set_iteration_maintenance_status(
    radicado: str,
    iteration: int,
    *,
    status: str,
    message: Optional[str] = None,
    run_id: Optional[str] = None,
) -> Dict[str, Any]:
    state = load_case_state(radicado)
    iteration_key = str(iteration)
    entry = (state.get("iterations") or {}).get(iteration_key)
    if not entry:
        raise CaseStateError(f"La iteración {iteration} no existe para el caso {radicado}.")

    now = utc_now_iso()
    maintenance = dict(entry.get("maintenance") or {})
    maintenance["status"] = status
    if status == "queued":
        maintenance["queued_at"] = now
        maintenance.pop("started_at", None)
        maintenance.pop("finished_at", None)
        maintenance.pop("run_id", None)
    elif status == "running":
        maintenance.setdefault("queued_at", now)
        maintenance["started_at"] = now
        maintenance.pop("finished_at", None)
    else:
        maintenance.setdefault("queued_at", now)
        maintenance.setdefault("started_at", now)
        maintenance["finished_at"] = now
    if message is not None:
        maintenance["message"] = message
    if run_id is not None:
        maintenance["run_id"] = run_id

    entry["maintenance"] = maintenance
    entry["updated_at"] = now
    state["iterations"][iteration_key] = entry
    save_case_state(state)
    return state


def refresh_case_maintenance_state(state: Dict[str, Any]) -> Dict[str, Any]:
    timeout_seconds = int(settings.OPENCLAW_MAINTENANCE_PENDING_TIMEOUT_SECONDS or 0)
    if timeout_seconds <= 0:
        return state

    now_dt = datetime.now(timezone.utc)
    changed = False
    for key, entry in (state.get("iterations") or {}).items():
        maintenance = dict(entry.get("maintenance") or {})
        status = (maintenance.get("status") or "").strip().lower()
        if status not in {"queued", "running"}:
            continue

        anchor = _parse_utc_iso(maintenance.get("started_at")) or _parse_utc_iso(maintenance.get("queued_at"))
        if not anchor:
            continue

        if (now_dt - anchor).total_seconds() < timeout_seconds:
            continue

        maintenance["status"] = "failed"
        maintenance["message"] = (
            "La actualización automática del backend excedió el tiempo máximo de espera. "
            "Puedes revisar logs o subir feedback nuevo cuando el backend vuelva a estar estable."
        )
        maintenance.setdefault("started_at", maintenance.get("queued_at") or utc_now_iso())
        maintenance["finished_at"] = utc_now_iso()
        entry["maintenance"] = maintenance
        entry["updated_at"] = utc_now_iso()
        state["iterations"][key] = entry
        changed = True

    if changed:
        save_case_state(state)
    return state


def finalize_generation(
    *,
    radicado: str,
    iteration: int,
    result: Dict[str, Any],
    comentario: str,
    template_id: Optional[str],
    staged_scans: Optional[List[Path]] = None,
    staged_docs: Optional[List[Path]] = None,
) -> Dict[str, Any]:
    state = ensure_case_state(radicado, comentario, template_id)

    if staged_scans is not None and staged_docs is not None:
        state["inputs"] = _copy_inputs_to_case(radicado, staged_scans, staged_docs)

    artifacts = _copy_generated_artifacts(radicado, iteration, result)
    change_report_rel = _write_iteration_change_report(
        radicado=radicado,
        state=state,
        iteration=iteration,
        artifacts=artifacts,
    )
    if change_report_rel:
        artifacts["change_report_path"] = change_report_rel
    now = utc_now_iso()
    state["latest_iteration"] = max(int(state.get("latest_iteration") or 0), iteration)
    state["status"] = "generated"
    state["iterations"][str(iteration)] = {
        "iteration": iteration,
        "status": "generated",
        "created_at": state["iterations"].get(str(iteration), {}).get("created_at") or now,
        "updated_at": now,
        "source_comentario": comentario or "(Sin comentarios)",
        "artifacts": artifacts,
        "feedback": None,
    }
    save_case_state(state)
    return state


def record_feedback(
    *,
    radicado: str,
    iteration: int,
    reviewed_docx_path: Path,
    comments_path: Path,
    comments: List[Dict[str, Any]],
    source_filename: Optional[str] = None,
    maintenance_status: str = "queued",
    maintenance_message: Optional[str] = None,
) -> Dict[str, Any]:
    state = load_case_state(radicado)
    iteration_key = str(iteration)
    if iteration_key not in state.get("iterations", {}):
        raise CaseStateError(f"La iteración {iteration} no existe para el caso {radicado}.")

    feedback = {
        "reviewed_docx_path": relative_to_case(radicado, reviewed_docx_path),
        "comments_json_path": relative_to_case(radicado, comments_path),
        "comments_count": len(comments),
        "uploaded_at": utc_now_iso(),
        "source_filename": source_filename or reviewed_docx_path.name,
    }
    state["status"] = "feedback_uploaded"
    state["iterations"][iteration_key]["status"] = "feedback_uploaded"
    state["iterations"][iteration_key]["updated_at"] = utc_now_iso()
    state["iterations"][iteration_key]["feedback"] = feedback
    state["iterations"][iteration_key]["maintenance"] = {
        "status": maintenance_status,
        "message": maintenance_message,
    }
    now = utc_now_iso()
    if maintenance_status == "queued":
        state["iterations"][iteration_key]["maintenance"]["queued_at"] = now
    elif maintenance_status == "running":
        state["iterations"][iteration_key]["maintenance"]["queued_at"] = now
        state["iterations"][iteration_key]["maintenance"]["started_at"] = now
    else:
        state["iterations"][iteration_key]["maintenance"]["queued_at"] = now
        state["iterations"][iteration_key]["maintenance"]["started_at"] = now
        state["iterations"][iteration_key]["maintenance"]["finished_at"] = now
    save_case_state(state)
    return state


def append_feedback_corpus_event(
    *,
    radicado: str,
    iteration: int,
    comments: List[Dict[str, Any]],
) -> None:
    state = load_case_state(radicado)
    iteration_key = str(iteration)
    entry = (state.get("iterations") or {}).get(iteration_key) or {}
    payload = {
        "timestamp": utc_now_iso(),
        "radicado": radicado,
        "iteration": iteration,
        "template_id": state.get("template_id"),
        "base_comentario": state.get("base_comentario"),
        "source_comentario": entry.get("source_comentario"),
        "comments_count": len(comments),
        "comments": comments,
        "inputs": state.get("inputs") or {},
        "artifacts": entry.get("artifacts") or {},
    }
    path = feedback_corpus_events_path()
    path.parent.mkdir(parents=True, exist_ok=True)
    with path.open("a", encoding="utf-8") as handle:
        handle.write(json.dumps(payload, ensure_ascii=False) + "\n")


def mark_iteration_in_progress(radicado: str) -> Dict[str, Any]:
    state = load_case_state(radicado)
    state["status"] = "iteration_in_progress"
    save_case_state(state)
    return state


def load_iteration_comments(radicado: str, iteration: int) -> List[Dict[str, Any]]:
    state = load_case_state(radicado)
    feedback = (state.get("iterations", {}).get(str(iteration), {}) or {}).get("feedback") or {}
    comments_rel = feedback.get("comments_json_path")
    if not comments_rel:
        return []
    path = resolve_case_path(radicado, comments_rel)
    if not path.exists():
        return []
    return json.loads(path.read_text(encoding="utf-8"))


def format_feedback_for_pipeline(state: Dict[str, Any]) -> str:
    chunks: List[str] = []
    for iteration_key in sorted(state.get("iterations", {}), key=lambda value: int(value)):
        iteration = state["iterations"][iteration_key]
        feedback = iteration.get("feedback") or {}
        comments_rel = feedback.get("comments_json_path")
        if not comments_rel:
            continue
        comments_path = resolve_case_path(str(state["radicado"]), comments_rel)
        if not comments_path.exists():
            continue
        comments = json.loads(comments_path.read_text(encoding="utf-8"))
        if not comments:
            continue
        chunks.append(f"ITERACION {iteration_key}: FEEDBACK WORD")
        for item in comments:
            anchor = (item.get("anchor_text") or "").strip()
            paragraph = (item.get("paragraph_text") or "").strip()
            comment_text = (item.get("comment_text") or "").strip()
            chunks.append(
                "\n".join(
                    [
                        f"- Comentario {item.get('comment_id')}: {comment_text}",
                        f"  Ancla: {anchor or '[sin ancla exacta]'}",
                        f"  Párrafo: {paragraph or '[sin párrafo detectado]'}",
                    ]
                )
            )
    return "\n\n".join(chunks).strip()


def compose_iteration_commentary(state: Dict[str, Any], extra_comment: Optional[str] = None) -> str:
    parts: List[str] = []
    base_comment = (state.get("base_comentario") or "").strip()
    if base_comment:
        parts.append(base_comment)
    if extra_comment and extra_comment.strip():
        parts.append(extra_comment.strip())

    feedback_block = format_feedback_for_pipeline(state)
    if feedback_block:
        parts.append(
            "APLICA OBLIGATORIAMENTE EL SIGUIENTE FEEDBACK DE REVISION WORD EN LA NUEVA ITERACION.\n"
            "NO IGNORES NINGUN COMENTARIO, NO RESUMAS, Y CORRIGE EL TEXTO LEGAL SEGUN ESTOS HALLAZGOS.\n\n"
            f"{feedback_block}"
        )

    return "\n\n".join(parts).strip() or "(Sin comentarios)"


def get_iteration_entry(state: Dict[str, Any], iteration: Optional[int] = None) -> Dict[str, Any]:
    latest = int(state.get("latest_iteration") or 0)
    resolved = iteration or latest
    entry = (state.get("iterations") or {}).get(str(resolved))
    if not entry:
        raise CaseStateError(f"La iteración {resolved} no existe para el caso {state['radicado']}.")
    return entry


def build_case_response(state: Dict[str, Any], iteration: Optional[int] = None) -> Dict[str, Any]:
    state = refresh_case_maintenance_state(state)
    radicado = str(state["radicado"])
    current = get_iteration_entry(state, iteration)
    current_iteration = int(current["iteration"])
    artifacts = current.get("artifacts") or {}

    iterations = []
    for key in sorted(state.get("iterations", {}), key=lambda value: int(value)):
        item = state["iterations"][key]
        item_artifacts = item.get("artifacts") or {}
        item_feedback = item.get("feedback") or {}
        iterations.append(
            {
                "iteration": int(item["iteration"]),
                "status": item.get("status"),
                "comments_count": int(item_feedback.get("comments_count") or 0),
                "maintenance_status": ((item.get("maintenance") or {}).get("status")),
                "artifacts": {
                    "docx_url": f"/cases/{radicado}/artifacts/docx?iteration={int(item['iteration'])}",
                    "pdf_url": f"/cases/{radicado}/artifacts/pdf?iteration={int(item['iteration'])}",
                    "change_report_url": (
                        f"/cases/{radicado}/artifacts/change-report?iteration={int(item['iteration'])}"
                        if item_artifacts.get("change_report_path")
                        else None
                    ),
                },
                "feedback_uploaded": bool(item_feedback),
            }
        )

    response = {
        "ok": True,
        "radicado": radicado,
        "current_iteration": current_iteration,
        "status": state.get("status"),
        "artifacts": {
            "docx_url": f"/cases/{radicado}/artifacts/docx?iteration={current_iteration}",
            "pdf_url": f"/cases/{radicado}/artifacts/pdf?iteration={current_iteration}",
            "change_report_url": (
                f"/cases/{radicado}/artifacts/change-report?iteration={current_iteration}"
                if artifacts.get("change_report_path")
                else None
            ),
        },
        "actions": {
            "case_url": f"/cases/{radicado}",
            "feedback_upload_url": f"/cases/{radicado}/feedback",
            "next_iteration_url": f"/cases/{radicado}/iterations/next",
        },
        "feedback": {
            "uploaded": bool(current.get("feedback")),
            "comments_count": int((current.get("feedback") or {}).get("comments_count") or 0),
        },
        "maintenance": current.get("maintenance") or None,
        "iterations": iterations,
    }

    if artifacts.get("docx_path"):
        response["docx_path"] = artifacts["docx_path"]
    if artifacts.get("pdf_path"):
        response["pdf_path"] = artifacts["pdf_path"]
    response["download_url"] = response["artifacts"]["pdf_url"]
    return response


def artifact_path_for_response(radicado: str, kind: str, iteration: Optional[int] = None) -> Path:
    state = load_case_state(radicado)
    entry = get_iteration_entry(state, iteration)
    relative_path = (entry.get("artifacts") or {}).get(f"{kind}_path")
    if not relative_path:
        raise CaseStateError(f"No existe artefacto {kind} para el caso {radicado}.")
    return resolve_case_path(radicado, relative_path)


@contextmanager
def case_lock(radicado: str) -> Iterator[None]:
    lock_path = case_dir(radicado) / ".iteration.lock"
    lock_path.parent.mkdir(parents=True, exist_ok=True)
    try:
        fd = os.open(lock_path, os.O_CREAT | os.O_EXCL | os.O_WRONLY)
    except FileExistsError as exc:
        raise CaseLockError(f"Ya hay una iteración en curso para el radicado {radicado}.") from exc
    try:
        os.write(fd, utc_now_iso().encode("utf-8"))
        yield
    finally:
        os.close(fd)
        if lock_path.exists():
            lock_path.unlink()
