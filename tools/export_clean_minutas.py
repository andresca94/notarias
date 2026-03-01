#!/usr/bin/env python3
"""
tools/export_clean_minutas.py
=============================
Convierte todas las minutas Word (.doc/.docx) de rag_store/actos/ a archivos
.txt limpios, eliminando:
  1. El formulario de calificación del encabezado (RADICADO, CLASE DE ACTO, etc.)
  2. El bloque OTORGAMIENTO Y AUTORIZACIÓN y todo lo que sigue (INSERTOS, UIAF, FIRMAS)

Los .txt resultantes son editables manualmente y tienen prioridad sobre los Word
en el sistema de RAG (rag_store.py los prefiere cuando existen).

Uso:
    python tools/export_clean_minutas.py [--rag-dir app/rag_store/actos]
    python tools/export_clean_minutas.py --dry-run   # muestra sin escribir
    python tools/export_clean_minutas.py --force      # sobreescribe .txt existentes
"""
from __future__ import annotations

import argparse
import re
import subprocess
import sys
import tempfile
from pathlib import Path


# ── Extracción ────────────────────────────────────────────────────────────────

def _extract_docx(path: Path) -> str:
    from docx import Document  # type: ignore
    doc = Document(str(path))
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text)
    return "\n".join(parts)


def _extract_doc(path: Path) -> str:
    import shutil

    def _via_libreoffice() -> str:
        with tempfile.TemporaryDirectory() as tmpdir:
            subprocess.run(
                ["libreoffice", "--headless", "--convert-to", "txt:Text",
                 "--outdir", tmpdir, str(path)],
                capture_output=True, timeout=30,
            )
            out = Path(tmpdir) / (path.stem + ".txt")
            return out.read_text(encoding="utf-8", errors="ignore") if out.exists() else ""

    def _via_textutil() -> str:
        r = subprocess.run(
            ["textutil", "-convert", "txt", "-stdout", str(path)],
            capture_output=True, timeout=30,
        )
        return r.stdout.decode("utf-8", errors="ignore") if r.returncode == 0 else ""

    if shutil.which("libreoffice"):
        return _via_libreoffice()
    if shutil.which("textutil"):
        return _via_textutil()
    raise RuntimeError("No se encontró libreoffice ni textutil para convertir .doc")


# ── Limpieza ──────────────────────────────────────────────────────────────────

def strip_formulario_header(text: str) -> str:
    """Elimina el formulario de calificación del inicio."""
    m = re.search(
        r'(Compareció\b|Comparecieron\b|En la ciudad de\b|En el municipio de\b)',
        text,
        re.IGNORECASE,
    )
    return text[m.start():].strip() if m else text


def strip_boilerplate_tail(text: str) -> str:
    """Elimina OTORGAMIENTO Y AUTORIZACIÓN y todo lo que sigue."""
    m = re.search(
        r'[-\s]{3,}O\s*T\s*O\s*R\s*G\s*A\s*M\s*I\s*E\s*N\s*T\s*O',
        text,
        re.IGNORECASE,
    )
    if m:
        return text[:m.start()].strip()
    # Fallback: cortar en INSERTOS
    m2 = re.search(r'^\s*INSERTOS\s*:', text, re.MULTILINE | re.IGNORECASE)
    if m2:
        return text[:m2.start()].strip()
    return text


def clean_excess_dashes(text: str) -> str:
    """Comprime líneas que son solo guiones (separadores decorativos)."""
    # Múltiples líneas de solo guiones → una sola línea de 78 guiones
    text = re.sub(r'(-{10,}\s*\n){2,}', '-' * 78 + '\n', text)
    # Guiones excesivos al final de párrafo → máximo 10
    text = re.sub(r'-{11,}$', '----------', text, flags=re.MULTILINE)
    return text


def process_file(path: Path, dry_run: bool = False, force: bool = False) -> str | None:
    """
    Procesa un archivo Word y retorna el texto limpio, o None si falla/omite.
    Guarda como .txt si no es dry_run.
    """
    out_path = path.with_suffix(".txt")
    if out_path.exists() and not force:
        print(f"  SKIP (ya existe, usa --force para sobreescribir): {out_path.name}")
        return None

    suffix = path.suffix.lower()
    try:
        if suffix == ".docx":
            raw = _extract_docx(path)
        elif suffix == ".doc":
            raw = _extract_doc(path)
        else:
            return None
    except Exception as e:
        print(f"  ERROR extrayendo {path.name}: {e}")
        return None

    if not raw.strip():
        print(f"  SKIP (vacío): {path.name}")
        return None

    cleaned = strip_formulario_header(raw)
    cleaned = strip_boilerplate_tail(cleaned)
    cleaned = clean_excess_dashes(cleaned)

    if not cleaned.strip():
        print(f"  SKIP (sin cuerpo luego de limpiar): {path.name}")
        return None

    if dry_run:
        print(f"\n{'='*78}")
        print(f"DRY-RUN: {path.name}")
        print(f"{'='*78}")
        print(cleaned[:2000])
        if len(cleaned) > 2000:
            print(f"... [{len(cleaned) - 2000} chars más] ...")
    else:
        out_path.write_text(cleaned, encoding="utf-8")
        print(f"  OK → {out_path.name}  ({len(cleaned)} chars)")

    return cleaned


# ── Main ──────────────────────────────────────────────────────────────────────

def main() -> None:
    parser = argparse.ArgumentParser(description=__doc__, formatter_class=argparse.RawDescriptionHelpFormatter)
    parser.add_argument("--rag-dir", default="app/rag_store/actos", help="Directorio de minutas Word")
    parser.add_argument("--dry-run", action="store_true", help="Mostrar resultado sin escribir archivos")
    parser.add_argument("--force", action="store_true", help="Sobreescribir .txt existentes")
    args = parser.parse_args()

    actos_dir = Path(args.rag_dir)
    if not actos_dir.is_dir():
        print(f"ERROR: directorio no existe: {actos_dir}", file=sys.stderr)
        sys.exit(1)

    word_files = sorted(
        p for p in actos_dir.iterdir()
        if p.suffix.lower() in (".doc", ".docx")
    )
    print(f"Procesando {len(word_files)} archivos Word en {actos_dir}\n")

    ok = skip = err = 0
    for path in word_files:
        print(f"→ {path.name}")
        result = process_file(path, dry_run=args.dry_run, force=args.force)
        if result is not None:
            ok += 1
        else:
            skip += 1

    print(f"\nResumen: {ok} generados, {skip} omitidos")
    if not args.dry_run:
        print(f"\nLos .txt están en {actos_dir}")
        print("Puedes editarlos manualmente. El sistema de RAG los usará con prioridad sobre los .doc/.docx")


if __name__ == "__main__":
    main()
