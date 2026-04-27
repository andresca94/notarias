from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document
from fastapi.testclient import TestClient

from app.core.config import settings
from app.main import app
from app.services.case_manager import feedback_corpus_events_path
from tests.helpers_docx import write_feedback_docx


@pytest.fixture
def client(tmp_path: Path, monkeypatch: pytest.MonkeyPatch):
    original_output_dir = settings.OUTPUT_DIR
    original_auto_tune_enabled = settings.OPENCLAW_AUTO_TUNE_ENABLED
    settings.OUTPUT_DIR = str(tmp_path / "outputs")
    settings.OPENCLAW_AUTO_TUNE_ENABLED = False

    call_counter = {"count": 0}

    async def fake_run_pipeline(scanner_paths, documentos_paths, comentario, template_id):
        call_counter["count"] += 1
        radicado = "26485"
        case_dir = Path(settings.OUTPUT_DIR) / f"CASE-{radicado}"
        debug_dir = case_dir / "debug"
        debug_dir.mkdir(parents=True, exist_ok=True)

        docx_path = case_dir / f"Minuta_Caso_{radicado}.docx"
        pdf_path = case_dir / f"Escritura_Caso_{radicado}.pdf"
        docx_path.write_bytes(f"draft-{call_counter['count']}".encode("utf-8"))
        pdf_path.write_bytes(f"pdf-{call_counter['count']}".encode("utf-8"))
        (debug_dir / "00_inputs_manifest.json").write_text("{}", encoding="utf-8")

        return {
            "radicado": radicado,
            "docx_path": str(docx_path),
            "pdf_path": str(pdf_path),
            "debug": {},
        }

    monkeypatch.setattr("app.api.routes.run_pipeline", fake_run_pipeline)

    with TestClient(app) as test_client:
        yield test_client, call_counter

    settings.OUTPUT_DIR = original_output_dir
    settings.OPENCLAW_AUTO_TUNE_ENABLED = original_auto_tune_enabled


def test_case_iteration_flow_preserves_previous_artifacts(client, tmp_path: Path):
    test_client, call_counter = client

    generate_response = test_client.post(
        "/notaria-v63-universal",
        data={"comentario": "Cliente solicita revisión"},
        files=[
            ("documentos", ("radicacion.pdf", b"%PDF-1.4", "application/pdf")),
            ("cedula", ("cc.pdf", b"%PDF-1.4", "application/pdf")),
        ],
    )
    assert generate_response.status_code == 200
    payload = generate_response.json()
    assert payload["current_iteration"] == 1
    assert payload["feedback"]["uploaded"] is False
    assert call_counter["count"] == 1

    feedback_path = write_feedback_docx(tmp_path / "feedback.docx")
    with feedback_path.open("rb") as feedback_file:
        feedback_response = test_client.post(
            "/cases/26485/feedback",
            files=[
                (
                    "feedback_docx",
                    ("feedback.docx", feedback_file.read(), "application/vnd.openxmlformats-officedocument.wordprocessingml.document"),
                )
            ],
        )
    assert feedback_response.status_code == 200
    feedback_payload = feedback_response.json()
    assert feedback_payload["feedback"]["uploaded"] is True
    assert feedback_payload["feedback"]["comments_count"] == 1

    next_response = test_client.post("/cases/26485/iterations/next")
    assert next_response.status_code == 200
    next_payload = next_response.json()
    assert next_payload["current_iteration"] == 2
    assert call_counter["count"] == 2

    iter1_docx = tmp_path / "outputs" / "CASE-26485" / "iterations" / "1" / "generated" / "Minuta_Caso_26485.docx"
    iter2_docx = tmp_path / "outputs" / "CASE-26485" / "iterations" / "2" / "generated" / "Minuta_Caso_26485.docx"
    assert iter1_docx.read_bytes() == b"draft-1"
    assert iter2_docx.read_bytes() == b"draft-2"


def test_next_iteration_requires_feedback(client):
    test_client, _ = client

    generate_response = test_client.post(
        "/notaria-v63-universal",
        files=[("documentos", ("radicacion.pdf", b"%PDF-1.4", "application/pdf"))],
    )
    assert generate_response.status_code == 200

    next_response = test_client.post("/cases/26485/iterations/next")
    assert next_response.status_code == 400
    assert "subir el DOCX revisado" in next_response.json()["detail"]


def test_feedback_upload_writes_corpus_event(client, tmp_path: Path):
    test_client, _ = client

    generate_response = test_client.post(
        "/notaria-v63-universal",
        files=[("documentos", ("radicacion.pdf", b"%PDF-1.4", "application/pdf"))],
    )
    assert generate_response.status_code == 200

    feedback_path = write_feedback_docx(tmp_path / "feedback.docx")
    with feedback_path.open("rb") as feedback_file:
        feedback_response = test_client.post(
            "/cases/26485/feedback",
            files=[
                (
                    "feedback_docx",
                    ("feedback.docx", feedback_file.read(), "application/vnd.openxmlformats-officedocument.wordprocessingml.document"),
                )
            ],
        )
    assert feedback_response.status_code == 200

    corpus_path = feedback_corpus_events_path()
    assert corpus_path.exists()
    payload = corpus_path.read_text(encoding="utf-8")
    assert '"radicado": "26485"' in payload


def test_feedback_upload_rejects_word_from_other_case(client, tmp_path: Path):
    test_client, _ = client

    generate_response = test_client.post(
        "/notaria-v63-universal",
        files=[("documentos", ("radicacion.pdf", b"%PDF-1.4", "application/pdf"))],
    )
    assert generate_response.status_code == 200

    feedback_path = write_feedback_docx(tmp_path / "Minuta_Caso_25963.docx")
    with feedback_path.open("rb") as feedback_file:
        feedback_response = test_client.post(
            "/cases/26485/feedback",
            files=[
                (
                    "feedback_docx",
                    ("Minuta_Caso_25963.docx", feedback_file.read(), "application/vnd.openxmlformats-officedocument.wordprocessingml.document"),
                )
            ],
        )
    assert feedback_response.status_code == 400
    assert "parece pertenecer a otro caso" in feedback_response.json()["detail"]


def test_next_iteration_exposes_change_report(
    client,
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
):
    test_client, _ = client

    call_counter = {"count": 0}

    async def fake_run_pipeline_with_docx(scanner_paths, documentos_paths, comentario, template_id):
        call_counter["count"] += 1
        radicado = "26485"
        case_dir = Path(settings.OUTPUT_DIR) / f"CASE-{radicado}"
        debug_dir = case_dir / "debug"
        debug_dir.mkdir(parents=True, exist_ok=True)

        docx_path = case_dir / f"Minuta_Caso_{radicado}.docx"
        pdf_path = case_dir / f"Escritura_Caso_{radicado}.pdf"

        doc = Document()
        doc.add_paragraph("Parrafo base de la iteracion inicial.")
        if call_counter["count"] > 1:
            doc.add_paragraph("Parrafo actualizado despues del feedback experto.")
        else:
            doc.add_paragraph("Parrafo anterior antes del feedback experto.")
        doc.save(docx_path)

        pdf_path.write_bytes(f"pdf-{call_counter['count']}".encode("utf-8"))
        (debug_dir / "00_inputs_manifest.json").write_text("{}", encoding="utf-8")

        return {
            "radicado": radicado,
            "docx_path": str(docx_path),
            "pdf_path": str(pdf_path),
            "debug": {},
        }

    monkeypatch.setattr("app.api.routes.run_pipeline", fake_run_pipeline_with_docx)

    generate_response = test_client.post(
        "/notaria-v63-universal",
        files=[("documentos", ("radicacion.pdf", b"%PDF-1.4", "application/pdf"))],
    )
    assert generate_response.status_code == 200

    feedback_path = write_feedback_docx(
        tmp_path / "feedback.docx",
        comment_text="Cambiar el parrafo principal segun la revision",
        paragraph_text="Parrafo anterior antes del feedback experto.",
        anchor_text="parrafo principal",
    )
    with feedback_path.open("rb") as feedback_file:
        feedback_response = test_client.post(
            "/cases/26485/feedback",
            files=[
                (
                    "feedback_docx",
                    ("feedback.docx", feedback_file.read(), "application/vnd.openxmlformats-officedocument.wordprocessingml.document"),
                )
            ],
        )
    assert feedback_response.status_code == 200

    next_response = test_client.post("/cases/26485/iterations/next")
    assert next_response.status_code == 200
    next_payload = next_response.json()

    report_url = next_payload["artifacts"]["change_report_url"]
    assert report_url == "/cases/26485/artifacts/change-report?iteration=2"

    report_response = test_client.get(report_url)
    assert report_response.status_code == 200
    report_text = report_response.text
    assert "Reporte de cambios de iteracion" in report_text
    assert "Cambiar el parrafo principal segun la revision" in report_text
    assert "Parrafo anterior antes del feedback experto." in report_text
    assert "Parrafo actualizado despues del feedback experto." in report_text


def test_feedback_upload_response_marks_maintenance_as_queued(client, tmp_path: Path):
    test_client, _ = client
    original_auto_tune_enabled = settings.OPENCLAW_AUTO_TUNE_ENABLED
    settings.OPENCLAW_AUTO_TUNE_ENABLED = True

    try:
        generate_response = test_client.post(
            "/notaria-v63-universal",
            files=[("documentos", ("radicacion.pdf", b"%PDF-1.4", "application/pdf"))],
        )
        assert generate_response.status_code == 200

        feedback_path = write_feedback_docx(tmp_path / "feedback.docx")
        with feedback_path.open("rb") as feedback_file:
            feedback_response = test_client.post(
                "/cases/26485/feedback",
                files=[
                    (
                        "feedback_docx",
                        ("feedback.docx", feedback_file.read(), "application/vnd.openxmlformats-officedocument.wordprocessingml.document"),
                    )
                ],
            )
        assert feedback_response.status_code == 200
        payload = feedback_response.json()
        assert payload["maintenance"]["status"] == "queued"
    finally:
        settings.OPENCLAW_AUTO_TUNE_ENABLED = original_auto_tune_enabled


def test_feedback_upload_marks_maintenance_as_skipped_when_auto_tune_disabled(client, tmp_path: Path):
    test_client, _ = client

    generate_response = test_client.post(
        "/notaria-v63-universal",
        files=[("documentos", ("radicacion.pdf", b"%PDF-1.4", "application/pdf"))],
    )
    assert generate_response.status_code == 200

    feedback_path = write_feedback_docx(tmp_path / "feedback.docx")
    with feedback_path.open("rb") as feedback_file:
        feedback_response = test_client.post(
            "/cases/26485/feedback",
            files=[
                (
                    "feedback_docx",
                    ("feedback.docx", feedback_file.read(), "application/vnd.openxmlformats-officedocument.wordprocessingml.document"),
                )
            ],
        )
    assert feedback_response.status_code == 200
    payload = feedback_response.json()
    assert payload["maintenance"]["status"] == "skipped"


def test_global_maintenance_blocks_new_generation_and_next_iteration(
    client,
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
):
    test_client, _ = client
    original_auto_tune_enabled = settings.OPENCLAW_AUTO_TUNE_ENABLED
    settings.OPENCLAW_AUTO_TUNE_ENABLED = True

    async def fake_run_auto_tune_for_feedback(**_kwargs):
        return None

    monkeypatch.setattr(
        "app.api.routes.run_auto_tune_for_feedback",
        fake_run_auto_tune_for_feedback,
    )

    try:
        generate_response = test_client.post(
            "/notaria-v63-universal",
            files=[("documentos", ("radicacion.pdf", b"%PDF-1.4", "application/pdf"))],
        )
        assert generate_response.status_code == 200

        feedback_path = write_feedback_docx(tmp_path / "feedback.docx")
        with feedback_path.open("rb") as feedback_file:
            feedback_response = test_client.post(
                "/cases/26485/feedback",
                files=[
                    (
                        "feedback_docx",
                        ("feedback.docx", feedback_file.read(), "application/vnd.openxmlformats-officedocument.wordprocessingml.document"),
                    )
                ],
            )
        assert feedback_response.status_code == 200

        blocked_generate = test_client.post(
            "/notaria-v63-universal",
            files=[("documentos", ("otro.pdf", b"%PDF-1.4", "application/pdf"))],
        )
        assert blocked_generate.status_code == 409
        assert "backend se está actualizando" in blocked_generate.json()["detail"]

        blocked_next = test_client.post("/cases/26485/iterations/next")
        assert blocked_next.status_code == 409
        assert "backend se está actualizando" in blocked_next.json()["detail"]
    finally:
        settings.OPENCLAW_AUTO_TUNE_ENABLED = original_auto_tune_enabled


def test_maintenance_callback_unlocks_next_iteration(
    client,
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
):
    test_client, _ = client
    original_auto_tune_enabled = settings.OPENCLAW_AUTO_TUNE_ENABLED
    original_admin_token = settings.INTERNAL_ADMIN_TOKEN
    settings.OPENCLAW_AUTO_TUNE_ENABLED = True
    settings.INTERNAL_ADMIN_TOKEN = "secret-token"

    async def fake_run_auto_tune_for_feedback(**_kwargs):
        return None

    monkeypatch.setattr(
        "app.api.routes.run_auto_tune_for_feedback",
        fake_run_auto_tune_for_feedback,
    )

    try:
        generate_response = test_client.post(
            "/notaria-v63-universal",
            files=[("documentos", ("radicacion.pdf", b"%PDF-1.4", "application/pdf"))],
        )
        assert generate_response.status_code == 200

        feedback_path = write_feedback_docx(tmp_path / "feedback.docx")
        with feedback_path.open("rb") as feedback_file:
            feedback_response = test_client.post(
                "/cases/26485/feedback",
                files=[
                    (
                        "feedback_docx",
                        ("feedback.docx", feedback_file.read(), "application/vnd.openxmlformats-officedocument.wordprocessingml.document"),
                    )
                ],
            )
        assert feedback_response.status_code == 200
        assert feedback_response.json()["maintenance"]["status"] == "queued"

        callback_response = test_client.post(
            "/admin/openclaw/backend-maintenance/status",
            headers={"x-admin-token": "secret-token"},
            json={
                "radicado": "26485",
                "iteration": 1,
                "status": "completed",
                "message": "Backend actualizado y verificado.",
            },
        )
        assert callback_response.status_code == 200
        assert callback_response.json()["maintenance"]["status"] == "completed"

        next_response = test_client.post("/cases/26485/iterations/next")
        assert next_response.status_code == 200
        assert next_response.json()["current_iteration"] == 2
    finally:
        settings.OPENCLAW_AUTO_TUNE_ENABLED = original_auto_tune_enabled
        settings.INTERNAL_ADMIN_TOKEN = original_admin_token
